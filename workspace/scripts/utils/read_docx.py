#!/usr/bin/env python3
# 读取 Word 文档内容

import sys
import os
from docx import Document

def read_docx(file_path):
    """读取 Word 文档内容"""
    try:
        doc = Document(file_path)
        content = []
        
        # 读取文档标题
        if doc.core_properties.title:
            content.append(f"文档标题: {doc.core_properties.title}")
        
        # 读取所有段落
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 只保留有内容的段落
                content.append(f"段落 {i+1}: {paragraph.text}")
        
        # 读取表格
        for table_idx, table in enumerate(doc.tables):
            content.append(f"\n表格 {table_idx+1}:")
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                content.append(f"  行 {row_idx+1}: {' | '.join(row_data)}")
        
        return "\n".join(content)
    
    except Exception as e:
        return f"读取文档时出错: {str(e)}"

if __name__ == "__main__":
    file_path = "/home/admin/Downloads/副本关于回流数据问题梳理.docx"
    
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        sys.exit(1)
    
    print(f"正在读取文档: {file_path}")
    print("=" * 80)
    
    content = read_docx(file_path)
    print(content)
    
    # 保存到文本文件以便分析
    output_path = "/home/admin/openclaw/workspace/回流数据问题分析.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    
    print(f"\n文档内容已保存到: {output_path}")