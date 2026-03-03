#!/usr/bin/env python3
import os
import sys
import docx
import re

def find_docx_files(search_term):
    """查找包含特定关键词的docx文件"""
    docx_files = []
    
    # 搜索常见目录
    search_dirs = [
        os.path.expanduser("~"),
        os.path.expanduser("~/Downloads"),
        os.path.expanduser("~/Documents"),
        os.path.expanduser("~/Desktop"),
    ]
    
    for search_dir in search_dirs:
        if not os.path.exists(search_dir):
            continue
            
        for root, dirs, files in os.walk(search_dir):
            for file in files:
                if file.lower().endswith('.docx'):
                    if search_term.lower() in file.lower():
                        full_path = os.path.join(root, file)
                        docx_files.append(full_path)
    
    return docx_files

def read_docx_content(file_path):
    """读取docx文件内容"""
    try:
        doc = docx.Document(file_path)
        content = []
        
        # 读取段落
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 读取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except Exception as e:
        return f"读取文件时出错: {str(e)}"

def main():
    # 搜索关键词
    search_terms = ["回流", "回流数据", "问题梳理"]
    
    print("🔍 正在搜索文档文件...")
    
    for term in search_terms:
        files = find_docx_files(term)
        if files:
            print(f"\n✅ 找到包含 '{term}' 的文档:")
            for file_path in files:
                print(f"   📄 {file_path}")
                
                # 读取文件内容
                print(f"\n📝 文件内容预览:")
                content = read_docx_content(file_path)
                print(content[:1000] + "..." if len(content) > 1000 else content)
                print("\n" + "="*80 + "\n")
            return
    
    # 如果没有找到特定文件，列出所有docx文件
    print("\n📋 所有找到的docx文件:")
    all_docx = []
    for root, dirs, files in os.walk(os.path.expanduser("~")):
        for file in files:
            if file.lower().endswith('.docx'):
                full_path = os.path.join(root, file)
                all_docx.append(full_path)
    
    for i, file_path in enumerate(all_docx[:20], 1):
        print(f"{i:2d}. {file_path}")
    
    if len(all_docx) > 20:
        print(f"... 还有 {len(all_docx) - 20} 个文件")

if __name__ == "__main__":
    main()