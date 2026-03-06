#!/usr/bin/env python3
"""
提取Word文档内容
"""

import docx
import sys
import os
from pathlib import Path

def extract_docx_content(file_path):
    """提取Word文档内容"""
    try:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f"错误：文件不存在 - {file_path}")
            return None
        
        print(f"正在提取文档：{file_path}")
        print("=" * 60)
        
        # 打开Word文档
        doc = docx.Document(file_path)
        
        # 提取所有段落
        content = []
        
        # 文档基本信息
        content.append(f"# 文档名称：{os.path.basename(file_path)}")
        content.append(f"文件大小：{os.path.getsize(file_path)} 字节")
        content.append(f"段落数量：{len(doc.paragraphs)}")
        content.append("")
        
        # 提取标题和内容
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # 只处理非空段落
                # 根据样式判断标题级别
                style = paragraph.style.name
                if 'Heading' in style:
                    # 提取标题级别
                    level = 1
                    if 'Heading' in style:
                        try:
                            level = int(style.replace('Heading ', ''))
                        except:
                            level = 1
                    
                    # 添加Markdown标题
                    content.append(f"{'#' * level} {text}")
                else:
                    content.append(text)
        
        return "\n".join(content)
        
    except Exception as e:
        print(f"提取文档时出错：{e}")
        return None

def main():
    # 文档路径
    docx_path = "/home/admin/Downloads/【校长日报】BI面板核数说明.docx"
    
    # 提取内容
    content = extract_docx_content(docx_path)
    
    if content:
        # 保存到文件
        output_path = "校长日报_BI面板核数说明_提取内容.md"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"\n✅ 文档内容已提取并保存到：{output_path}")
        print(f"📄 总字符数：{len(content)}")
        
        # 显示前500个字符预览
        print("\n📋 内容预览（前500字符）：")
        print("=" * 60)
        print(content[:500] + "..." if len(content) > 500 else content)
        print("=" * 60)
        
        return output_path
    else:
        print("❌ 提取文档内容失败")
        return None

if __name__ == "__main__":
    main()